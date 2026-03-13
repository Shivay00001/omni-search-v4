from fastapi import FastAPI, HTTPException, Query, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import Session
from sqlalchemy import select, func
import asyncio
import json
import datetime
import io
import re

from search_engine import search_web
from scraper import scrape_multiple
from sentinel_provider import sentinel
from insight_extractor import insight_extractor
from translator_manager import translator_manager
from database import get_db, init_db, AsyncSessionLocal
from models import Conversation, Message, ResearchStep, RL_Feedback
from research_agent import research_agent

app = FastAPI(title="Omni Search & Summary Engine - V4")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    await init_db()

from typing import List, Dict, Optional

# Focus mode search prefixes
FOCUS_MODE_PREFIXES = {
    "academic": "site:arxiv.org OR site:scholar.google.com OR site:pubmed.ncbi.nlm.nih.gov ",
    "news": "latest news ",
    "code": "site:github.com OR site:stackoverflow.com OR site:dev.to ",
    "general": ""
}

class SearchRequest(BaseModel):
    query: str
    stream: bool = True
    target_lang: str = "en"
    conversation_id: Optional[int] = None
    focus_mode: str = "general"  # academic, news, code, general

class FeedbackRequest(BaseModel):
    query: str
    url: str
    score: float # 1.0 for like, -1.0 for dislike

class RelatedQuestionsRequest(BaseModel):
    query: str
    context: str = ""

@app.get("/")
async def root():
    return {"status": "online", "version": "4.0.0-ultimate"}

@app.get("/languages")
async def get_languages():
    """Returns all supported languages with metadata."""
    return translator_manager.get_supported_languages()

@app.get("/conversations")
async def list_conversations(db: Session = Depends(get_db)):
    result = await db.execute(select(Conversation).order_by(Conversation.created_at.desc()))
    return result.scalars().all()

@app.get("/conversations/{conv_id}")
async def get_conversation(conv_id: int, db: Session = Depends(get_db)):
    conv_result = await db.execute(select(Conversation).filter(Conversation.id == conv_id))
    conv = conv_result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    msg_result = await db.execute(select(Message).filter(Message.conversation_id == conv_id).order_by(Message.created_at.asc()))
    messages = msg_result.scalars().all()
    return {"conversation": conv, "messages": messages}

@app.post("/feedback")
async def save_feedback(req: FeedbackRequest, db: Session = Depends(get_db)):
    from urllib.parse import urlparse
    domain = urlparse(req.url).netloc
    feedback = RL_Feedback(query=req.query, url=req.url, domain=domain, score=req.score)
    db.add(feedback)
    await db.commit()
    return {"status": "recorded"}

@app.get("/discover")
async def get_discover():
    """Provides dynamic trending research topics from live sources."""
    from trend_engine import trend_engine
    trends = await trend_engine.get_trending_topics()
    return trends

@app.post("/related_questions")
async def get_related_questions(req: RelatedQuestionsRequest):
    """Generates 3 follow-up questions using LLM."""
    prompt = f"""Based on this research query and context, generate exactly 3 concise follow-up questions a user might ask next.

Query: {req.query}
Context summary: {req.context[:500] if req.context else 'No prior context'}

Output ONLY a JSON array of 3 strings. Example: ["Question 1?", "Question 2?", "Question 3?"]"""
    
    try:
        result = await sentinel.generate_full(prompt)
        start = result.find('[')
        end = result.rfind(']') + 1
        if start != -1 and end > start:
            questions = json.loads(result[start:end])
            return questions[:3]
    except Exception as e:
        print(f"Related questions error: {e}")
    
    return [
        f"What are the practical applications of {req.query}?",
        f"How does {req.query} compare to alternatives?",
        f"What are the latest developments in {req.query}?"
    ]

def _build_markdown_export(conversation, messages):
    """Build Markdown export string."""
    content = f"# Research Session: {conversation.title}\n"
    content += f"*Exported on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n---\n\n"
    for msg in messages:
        if msg.role == "user":
            content += f"## 🔍 Query\n{msg.content}\n\n"
        else:
            content += f"## 📝 Research Response\n{msg.content}\n\n---\n\n"
    return content

@app.get("/conversations/{conv_id}/export")
async def export_conversation(conv_id: int, format: str = "md", db: AsyncSession = Depends(get_db)):
    """Exports a conversation in MD, PDF, or DOCX format."""
    conversation = await db.get(Conversation, conv_id)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    stmt = select(Message).where(Message.conversation_id == conv_id).order_by(Message.id)
    result = await db.execute(stmt)
    messages = result.scalars().all()
    
    title = conversation.title or f"Research Session #{conv_id}"
    
    if format == "md":
        content = _build_markdown_export(conversation, messages)
        return Response(
            content=content.encode('utf-8'),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="research_{conv_id}.md"'}
        )
    
    elif format == "pdf":
        try:
            from fpdf import FPDF
            
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            
            # Title
            pdf.set_font("Helvetica", "B", 18)
            pdf.cell(0, 12, title.encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(0, 8, f"Exported on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(8)
            
            for msg in messages:
                if msg.role == "user":
                    pdf.set_font("Helvetica", "B", 13)
                    pdf.set_text_color(59, 130, 246)
                    pdf.cell(0, 10, "Query", ln=True)
                    pdf.set_text_color(0, 0, 0)
                    pdf.set_font("Helvetica", "", 11)
                    safe_text = msg.content.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 7, safe_text)
                    pdf.ln(4)
                else:
                    pdf.set_font("Helvetica", "B", 13)
                    pdf.set_text_color(16, 185, 129)
                    pdf.cell(0, 10, "Research Response", ln=True)
                    pdf.set_text_color(0, 0, 0)
                    pdf.set_font("Helvetica", "", 10)
                    # Strip markdown formatting for PDF
                    clean = re.sub(r'[#*_`]', '', msg.content or '')
                    safe_text = clean.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 6, safe_text)
                    pdf.ln(6)
                    pdf.set_draw_color(200, 200, 200)
                    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                    pdf.ln(4)
            
            pdf_bytes = pdf.output()
            return Response(
                content=bytes(pdf_bytes),
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="research_{conv_id}.pdf"'}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="fpdf2 not installed. Run: pip install fpdf2")
    
    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title_para = doc.add_heading(title, level=1)
            subtitle = doc.add_paragraph(f"Exported on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
            subtitle.style.font.size = Pt(10)
            subtitle.style.font.color.rgb = RGBColor(128, 128, 128)
            
            for msg in messages:
                if msg.role == "user":
                    heading = doc.add_heading("Query", level=2)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(59, 130, 246)
                    doc.add_paragraph(msg.content)
                else:
                    heading = doc.add_heading("Research Response", level=2)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(16, 185, 129)
                    # Split into paragraphs
                    clean = re.sub(r'[#*_`]', '', msg.content or '')
                    for para_text in clean.split('\n'):
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
                    doc.add_paragraph("─" * 50)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="research_{conv_id}.docx"'}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}. Use md, pdf, or docx.")

@app.post("/search")
async def search(request: SearchRequest, db: Session = Depends(get_db)):
    if not request.query:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    
    # Translation Layer (Input)
    original_query = request.query
    if request.target_lang != "en":
        request.query = await translator_manager.translate_to_en(request.query)

    # Apply focus mode prefix
    focus_prefix = FOCUS_MODE_PREFIXES.get(request.focus_mode, "")
    search_query = focus_prefix + request.query if focus_prefix else request.query

    async def stream_generator():
        try:
            nonlocal request, search_query
            conv_id = request.conversation_id
            
            # Initialize/Save Session
            async with AsyncSessionLocal() as session:
                if not conv_id:
                    new_conv = Conversation(title=original_query[:80])
                    session.add(new_conv)
                    await session.flush()
                    conv_id = new_conv.id
                    await session.commit()
                    yield f"__CONV_ID__:{conv_id}\n"
                
                user_msg = Message(conversation_id=conv_id, role="user", content=original_query)
                session.add(user_msg)
                await session.commit()

            full_response = ""
            context_data = []
            is_task = any(word in original_query.lower() for word in ["find and", "buy", "compare", "action", "task", "conduct", "research for"])
            
            if is_task:
                async for step in research_agent.execute_agent_task(search_query):
                    if isinstance(step, str):
                        yield f"__THOUGHT__:{step}\n"
                    else:
                        context_data.extend(step)
            else:
                async for step in research_agent.conduct_linear_research(search_query):
                    if isinstance(step, str):
                        yield f"__THOUGHT__:{step}\n"
                    else:
                        context_data.extend(step)

            # Stream Sources for transparency
            source_list = []
            for i, r in enumerate(context_data):
                source_list.append({
                    "index": i + 1,
                    "title": r.get("title", "Unknown"),
                    "url": r.get("url", "#"),
                    "snippet": r.get("snippet", "")[:200],
                    "score": round(r.get("score", 0), 2)
                })
            yield f"__SOURCES__:{json.dumps(source_list)}\n"

            # Insight Extraction
            all_text = " ".join([r.get("snippet", "") for r in context_data])
            insights = insight_extractor.extract_facts(all_text)
            yield f"__INSIGHTS__:{json.dumps(insights)}\n"
            
            # Build citation-aware context for LLM
            citation_lines = []
            for i, r in enumerate(context_data):
                citation_lines.append(f"[{i+1}] Source: {r['url']}\nTitle: {r.get('title', '')}\nSnippet: {r['snippet']}")
            context_text = "\n\n".join(citation_lines)[:6000]
            
            summary_prompt = f"""You are Omni, a premium AI Search Intelligence Engine.

Topic: {request.query}

Research Sources:
{context_text}

INSTRUCTIONS:
- Provide a comprehensive, well-structured answer using the sources above.
- Use inline citations like [1], [2], [3] to reference specific sources by their number.
- Use Markdown formatting (bold, headers, bullet points).
- Be factual and cite every major claim.
- End with a brief "Key Takeaway" section.
"""
            
            async for chunk in sentinel.generate_stream_response([{"role": "user", "content": summary_prompt}]):
                chunk_str = chunk
                if request.target_lang != "en":
                    chunk_str = await translator_manager.translate_to_target(chunk, request.target_lang)
                if chunk_str is not None:
                    yield chunk_str
                    full_response += chunk_str

            # Save Assistant Response
            async with AsyncSessionLocal() as session:
                assistant_msg = Message(conversation_id=conv_id, role="assistant", content=full_response)
                session.add(assistant_msg)
                await session.commit()
        except Exception as e:
            yield f"__ERROR__: {str(e)}\n"
            print(f"Stream generation error: {e}")

    return StreamingResponse(stream_generator(), media_type="text/plain")

@app.get("/test_search")
async def test_search(q: str = Query(...)):
    return await search_web(q)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8089)
